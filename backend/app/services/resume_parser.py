# =============================================
# 简历文件解析器 — PDF / Word 文本提取
# =============================================
# 支持 .pdf 和 .docx 格式
# 提取纯文本内容，用于 AI 分析或导入到编辑器
# =============================================

import io
from typing import Optional


async def parse_pdf(file_bytes: bytes) -> str:
    """从 PDF 文件提取纯文本"""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    texts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            texts.append(text.strip())
    return "\n\n".join(texts)


async def parse_docx(file_bytes: bytes) -> str:
    """从 Word (.docx) 文件提取纯文本"""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    texts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            texts.append(text)
    # 也提取表格中的文本（简历常用表格布局）
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                texts.append(" | ".join(cells))
    return "\n".join(texts)


async def parse_resume_file(filename: str, file_bytes: bytes) -> Optional[str]:
    """
    根据文件扩展名自动选择解析器

    返回: 提取的纯文本，格式不支持返回 None
    """
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return await parse_pdf(file_bytes)
    elif lower.endswith(".docx"):
        return await parse_docx(file_bytes)
    return None
